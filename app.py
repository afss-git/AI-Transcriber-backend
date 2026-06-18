"""
AI Transcriber — free audio/video transcription + AI document assistant.

Powered by faster-whisper (Whisper) for transcription and Groq (Llama 3) for
AI formatting/summarisation. Disconnect-proof: saves each line to disk live.
"""

import os
import re
import glob
import time
import threading
import datetime
import gradio as gr
from faster_whisper import WhisperModel

try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

try:
    from docx import Document as DocxDoc
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL = "llama-3.3-70b-versatile"

APP_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(APP_DIR, "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Model handling
# ---------------------------------------------------------------------------
MODEL_CACHE = {}
MODEL_INFO = {
    "tiny":     "Fastest, lowest accuracy. Good for quick drafts.",
    "base":     "Very fast, basic accuracy.",
    "small":    "Fast, decent accuracy. Good balance on a CPU.",
    "medium":   "Slower, good accuracy. Solid for clear speech (slow on CPU).",
    "large-v3": "Slowest on CPU, best accuracy. Use for names, accents, foreign terms.",
}


def get_model(size: str) -> WhisperModel:
    if size not in MODEL_CACHE:
        MODEL_CACHE[size] = WhisperModel(size, device="cpu", compute_type="int8", cpu_threads=4)
    return MODEL_CACHE[size]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _srt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _safe_base(name: str) -> str:
    base = os.path.splitext(os.path.basename(name))[0]
    cleaned = "".join(c for c in base if c.isalnum() or c in " ._-").strip()
    return cleaned or "audio"


LINE_RE = re.compile(r"\[(\d+(?:\.\d+)?)s -> (\d+(?:\.\d+)?)s\]\s?(.*)")


def _plain_text(raw: str) -> str:
    parts = []
    for line in raw.splitlines():
        m = LINE_RE.match(line.strip())
        parts.append(m.group(3) if m else line.strip())
    return " ".join(p for p in parts if p)


# ---------------------------------------------------------------------------
# Background transcription worker
# ---------------------------------------------------------------------------
JOBS = {}


def _worker(job_id, audio_file, model_size, lang, task, multilingual, prompt):
    job = JOBS[job_id]
    try:
        model = get_model(model_size)
        segments, info = model.transcribe(
            audio_file,
            language=lang,
            task=task,
            multilingual=multilingual,
            initial_prompt=prompt,
            vad_filter=True,
        )
        total = getattr(info, "duration", 0) or 0
        job["total"] = total
        job["lang"] = info.language
        with open(job["txt_path"], "w", encoding="utf-8") as f:
            for seg in segments:
                f.write(f"[{seg.start:.1f}s -> {seg.end:.1f}s] {seg.text.strip()}\n")
                f.flush()
                job["done"] = seg.end
                job["progress"] = min(seg.end / total, 0.999) if total else 0.0
        job["progress"] = 1.0
        job["status"] = "completed"
    except Exception as e:
        job["status"] = "error"
        job["error"] = str(e)


def start_transcription(audio_file, model_size, language, task_mode, multilingual, key_terms):
    if not audio_file:
        yield "⚠️ Please upload an audio or video file first.", "", None, ""
        return

    multilingual = bool(multilingual)
    lang = None if (language == "auto-detect" or multilingual) else language
    task = "translate" if str(task_mode).startswith("Translate") else "transcribe"
    prompt = key_terms.strip() if key_terms else None

    ts = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    job_id = f"{_safe_base(audio_file)}_{ts}"
    txt_path = os.path.join(OUTPUT_DIR, f"{job_id}.txt")
    JOBS[job_id] = {"status": "running", "progress": 0.0, "done": 0,
                    "total": 0, "txt_path": txt_path, "error": None, "lang": None}

    threading.Thread(
        target=_worker,
        args=(job_id, audio_file, model_size, lang, task, multilingual, prompt),
        daemon=True,
    ).start()

    yield f"🟡 Starting… loading '{model_size}' model (first run downloads it).\n💾 Saving to: {txt_path}", "", None, ""

    while True:
        job = JOBS[job_id]
        text = ""
        if os.path.exists(txt_path):
            try:
                with open(txt_path, encoding="utf-8") as f:
                    text = f.read()
            except Exception:
                pass
        dl = txt_path if os.path.exists(txt_path) else None

        if job["status"] == "completed":
            yield f"✅ Done! Saved to: {txt_path}", text, dl, text
            return
        if job["status"] == "error":
            yield f"❌ Error: {job['error']}\n(Partial transcript saved — use Recover tab.)", text, dl, text
            return

        pct = job["progress"] * 100
        total = job["total"]
        if total:
            status = (f"🔴 Transcribing — {pct:.0f}%  ({job['done']:.0f}s / {total:.0f}s)  |  lang: {job['lang']}\n"
                      f"💾 Saving live — safe to close and recover later.")
        else:
            status = f"🔴 Decoding audio & warming up…\n💾 Saving to: {txt_path}"
        yield status, text, dl, text
        time.sleep(2)


# ---------------------------------------------------------------------------
# Recovery / export helpers
# ---------------------------------------------------------------------------
def list_saved():
    files = sorted(glob.glob(os.path.join(OUTPUT_DIR, "*.txt")), key=os.path.getmtime, reverse=True)
    return gr.update(choices=[os.path.basename(f) for f in files])


def load_saved(name):
    if not name:
        return "Pick a saved transcript first.", None
    path = os.path.join(OUTPUT_DIR, name)
    if not os.path.exists(path):
        return "File not found (refresh the list).", None
    with open(path, encoding="utf-8") as f:
        return f.read(), path


def export_as(name, fmt):
    if not name:
        return None
    path = os.path.join(OUTPUT_DIR, name)
    if not os.path.exists(path):
        return None
    rows = []
    with open(path, encoding="utf-8") as f:
        for line in f:
            m = LINE_RE.match(line.strip())
            if m:
                rows.append((float(m.group(1)), float(m.group(2)), m.group(3)))
    stem = os.path.splitext(path)[0]
    if fmt == "Plain text (.txt)":
        out = stem + "_plain.txt"
        with open(out, "w", encoding="utf-8") as f:
            f.write(" ".join(r[2] for r in rows))
    else:
        out = stem + ".srt"
        with open(out, "w", encoding="utf-8") as f:
            for i, (s, e, t) in enumerate(rows, start=1):
                f.write(f"{i}\n{_srt_time(s)} --> {_srt_time(e)}\n{t}\n\n")
    return out


# ---------------------------------------------------------------------------
# AI Assistant
# ---------------------------------------------------------------------------
QUICK_PROMPTS = {
    "📋 Meeting Minutes":    "Format this transcript as professional meeting minutes. Include: Date (if mentioned), Attendees (if mentioned), Agenda items discussed, Key decisions made, Action items with owners, and Next steps.",
    "📝 Executive Summary":  "Write a concise executive summary in 3–5 paragraphs covering: main topic, key points discussed, conclusions reached, and any recommendations.",
    "✅ Action Items":        "Extract every action item, task, and decision from this transcript. Present as a numbered list. For each item include: the task, person responsible (if mentioned), and deadline (if mentioned).",
    "📚 Lecture Notes":      "Convert this lecture transcript into well-structured study notes. Include: main topic, key concepts with explanations, important definitions, examples given, and a brief summary.",
    "📄 Professional Report":"Rewrite this transcript as a formal professional report with these sections: Executive Summary, Background/Context, Main Discussion Points, Conclusions, and Recommendations.",
}


def chat_with_transcript(user_msg, history, transcript):
    def _reply(msg, hist, text):
        return hist + [{"role": "user", "content": user_msg}, {"role": "assistant", "content": text}]

    if not GROQ_AVAILABLE:
        h = _reply(user_msg, history, "❌ Groq not installed. Add `groq` to requirements.txt and redeploy.")
        return h, h

    if not GROQ_API_KEY:
        h = _reply(user_msg, history, "❌ No GROQ_API_KEY found. Add it as a Space secret (Settings → Variables and secrets).")
        return h, h

    if not transcript or not transcript.strip():
        h = _reply(user_msg, history, "⚠️ No transcript loaded. Paste one or transcribe audio first, then click 🔄 Load Transcript.")
        return h, h

    plain = _plain_text(transcript)
    system = (
        "You are a professional document assistant. The user has transcribed an audio recording "
        "and wants help processing it into a useful document.\n\n"
        f"FULL TRANSCRIPT:\n---\n{plain[:12000]}\n---\n\n"
        "Help the user with whatever they ask. Be clear, professional, and well-structured. "
        "Use markdown formatting: ## for headings, **bold** for key terms, - for bullet lists."
    )

    api_messages = [{"role": "system", "content": system}]
    for msg in history:
        api_messages.append({"role": msg["role"], "content": msg["content"]})
    api_messages.append({"role": "user", "content": user_msg})

    try:
        client = Groq(api_key=GROQ_API_KEY)
        resp = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=api_messages,
            max_tokens=4096,
            temperature=0.3,
        )
        reply = resp.choices[0].message.content
    except Exception as e:
        reply = f"❌ Groq API error: {e}"

    new_history = history + [{"role": "user", "content": user_msg}, {"role": "assistant", "content": reply}]
    return new_history, new_history


def generate_docx(history):
    if not DOCX_AVAILABLE or not history:
        return None
    last_reply = ""
    for msg in reversed(history):
        if msg.get("role") == "assistant":
            last_reply = msg.get("content", "")
            break
    if not last_reply:
        return None

    doc = DocxDoc()
    doc.add_heading("AI Document", 0)
    for line in last_reply.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
        else:
            p = doc.add_paragraph()
            parts = re.split(r"\*\*(.*?)\*\*", line)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
    ts = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    out = os.path.join(OUTPUT_DIR, f"document_{ts}.docx")
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------
with gr.Blocks(title="AI Transcriber") as demo:
    transcript_state = gr.State("")
    chat_state = gr.State([])

    gr.Markdown(
        "# 🎙️ AI Transcriber\n"
        "**Step 1:** Transcribe your audio/video &nbsp;→&nbsp; "
        "**Step 2:** Use the AI Assistant to format it into a professional document."
    )

    with gr.Tabs():

        # ── TAB 1: TRANSCRIBE ────────────────────────────────────────────────
        with gr.Tab("🎙️ Transcribe"):
            with gr.Row():
                with gr.Column(scale=1):
                    audio_in = gr.Audio(label="Upload audio / video", type="filepath")
                    file_in = gr.File(
                        label="…or drop any file here (mp4, mp3, wav, m4a, webm, …)",
                        file_types=["audio", "video"],
                    )
                    model_dd = gr.Dropdown(
                        choices=list(MODEL_INFO.keys()), value="small",
                        label="Model (accuracy vs. speed)",
                    )
                    model_note = gr.Markdown(f"*{MODEL_INFO['small']}*")
                    lang_dd = gr.Dropdown(
                        choices=["auto-detect", "en", "ar", "fr", "es", "de",
                                 "ur", "hi", "tr", "id", "ha"],
                        value="auto-detect", label="Language",
                    )
                    multi_cb = gr.Checkbox(
                        value=False, label="🌐 Multilingual / code-switching",
                        info="Turn ON for mixed Arabic+English or similar audio.",
                    )
                    task_dd = gr.Radio(
                        choices=["Transcribe (keep original languages)",
                                 "Translate everything to English"],
                        value="Transcribe (keep original languages)", label="Task",
                    )
                    terms_tb = gr.Textbox(
                        label="Key names & terms (optional)",
                        placeholder="e.g. Hamza Yusuf, taqwa, Qur'an, dunya",
                        lines=2,
                        info="Biases the model toward correct spelling of names/jargon.",
                    )
                    go_btn = gr.Button("Transcribe ▶", variant="primary", size="lg")

                with gr.Column(scale=2):
                    status_md = gr.Markdown("Idle. Upload a file and click **Transcribe ▶**.")
                    out_text = gr.Textbox(label="Live transcript", lines=20)
                    out_file = gr.File(label="Download .txt")
                    gr.Markdown(
                        "✅ *Transcription complete? Switch to the **🤖 AI Assistant** tab "
                        "and click **🔄 Load Transcript**.*"
                    )

        # ── TAB 2: AI ASSISTANT ──────────────────────────────────────────────
        with gr.Tab("🤖 AI Assistant"):
            gr.Markdown(
                "### Format, summarise, or restructure your transcript with AI.\n"
                "Click **🔄 Load Transcript** first, then choose a quick action or type your own instruction."
            )
            with gr.Row():
                with gr.Column(scale=1):
                    gr.Markdown("**⚡ Quick Actions**")
                    btn_meeting = gr.Button("📋 Meeting Minutes", size="sm")
                    btn_summary = gr.Button("📝 Executive Summary", size="sm")
                    btn_actions = gr.Button("✅ Action Items", size="sm")
                    btn_lecture = gr.Button("📚 Lecture Notes", size="sm")
                    btn_report  = gr.Button("📄 Professional Report", size="sm")
                    gr.Markdown("---")
                    load_t_btn = gr.Button("🔄 Load from Transcribe tab", variant="secondary")
                    gr.Markdown("*— or paste your own transcript below —*")
                    paste_box = gr.Textbox(
                        label="Paste transcript here",
                        lines=5,
                        placeholder="Paste any transcript text here, then click Use This Transcript ↓",
                    )
                    use_paste_btn = gr.Button("✅ Use This Transcript", variant="primary")
                    transcript_preview = gr.Textbox(
                        label="Active transcript (preview)",
                        lines=5, interactive=False,
                        placeholder="Load or paste a transcript to get started.",
                    )

                with gr.Column(scale=2):
                    chatbot = gr.Chatbot(label="AI Assistant", height=420)
                    user_input = gr.Textbox(
                        label="Your instruction",
                        placeholder="e.g. Write this up as formal meeting minutes for our board…",
                        lines=3,
                    )
                    with gr.Row():
                        send_btn  = gr.Button("Send ▶", variant="primary", scale=3)
                        clear_btn = gr.Button("🗑️ Clear chat", scale=1)
                    gr.Markdown("**Download last AI reply as:**")
                    with gr.Row():
                        docx_btn  = gr.Button("📄 Word Document (.docx)", variant="secondary")
                        docx_file = gr.File(label="Word file ready to download")

        # ── TAB 3: RECOVER ───────────────────────────────────────────────────
        with gr.Tab("🛟 Recover / Export"):
            gr.Markdown(
                "Reopen any saved transcript — even partial ones from interrupted sessions."
            )
            with gr.Row():
                saved_dd    = gr.Dropdown(choices=[], label="Saved transcripts (newest first)", scale=3)
                refresh_btn = gr.Button("🔄 Refresh", scale=1)
                load_btn    = gr.Button("📂 Load", variant="primary", scale=1)
            rec_text = gr.Textbox(label="Transcript", lines=15)
            rec_file = gr.File(label="Download")
            with gr.Row():
                exp_fmt = gr.Radio(
                    choices=["Plain text (.txt)", "Subtitles (.srt)"],
                    value="Plain text (.txt)", label="Export format", scale=2,
                )
                exp_btn  = gr.Button("⬇️ Export", scale=1)
                exp_file = gr.File(label="Exported file", scale=2)

    # ── Wiring ───────────────────────────────────────────────────────────────
    file_in.change(lambda f: f, inputs=file_in, outputs=audio_in)
    model_dd.change(lambda m: f"*{MODEL_INFO[m]}*", inputs=model_dd, outputs=model_note)

    go_btn.click(
        start_transcription,
        inputs=[audio_in, model_dd, lang_dd, task_dd, multi_cb, terms_tb],
        outputs=[status_md, out_text, out_file, transcript_state],
    )

    load_t_btn.click(
        fn=lambda t: (t[:3000] + "\n…(truncated for preview)" if len(t) > 3000 else t),
        inputs=transcript_state,
        outputs=transcript_preview,
    )

    use_paste_btn.click(
        fn=lambda p: (p, p[:3000] + "\n…(truncated for preview)" if len(p) > 3000 else p),
        inputs=paste_box,
        outputs=[transcript_state, transcript_preview],
    )

    btn_meeting.click(fn=lambda: QUICK_PROMPTS["📋 Meeting Minutes"],    outputs=user_input)
    btn_summary.click(fn=lambda: QUICK_PROMPTS["📝 Executive Summary"],  outputs=user_input)
    btn_actions.click(fn=lambda: QUICK_PROMPTS["✅ Action Items"],        outputs=user_input)
    btn_lecture.click(fn=lambda: QUICK_PROMPTS["📚 Lecture Notes"],      outputs=user_input)
    btn_report.click( fn=lambda: QUICK_PROMPTS["📄 Professional Report"], outputs=user_input)

    send_btn.click(
        fn=chat_with_transcript,
        inputs=[user_input, chat_state, transcript_state],
        outputs=[chatbot, chat_state],
    ).then(fn=lambda: "", outputs=user_input)

    user_input.submit(
        fn=chat_with_transcript,
        inputs=[user_input, chat_state, transcript_state],
        outputs=[chatbot, chat_state],
    ).then(fn=lambda: "", outputs=user_input)

    clear_btn.click(fn=lambda: ([], []), outputs=[chatbot, chat_state])

    docx_btn.click(fn=generate_docx, inputs=chat_state, outputs=docx_file)

    demo.load(list_saved, outputs=saved_dd)
    refresh_btn.click(list_saved, outputs=saved_dd)
    load_btn.click(load_saved, inputs=saved_dd, outputs=[rec_text, rec_file])
    exp_btn.click(export_as, inputs=[saved_dd, exp_fmt], outputs=exp_file)


if __name__ == "__main__":
    demo.launch(inbrowser=True, theme=gr.themes.Soft())
