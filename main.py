"""
AI Transcriber — Flask backend
Render: build=pip install -r requirements.txt | start=gunicorn main:app --bind 0.0.0.0:$PORT --workers 1 --timeout 120
"""

import os
import re
import uuid
import json
import threading
import datetime
from pathlib import Path

from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS

from faster_whisper import WhisperModel

try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

try:
    from docx import Document as DocxDoc
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL   = "llama-3.3-70b-versatile"

APP_DIR    = Path(__file__).parent
UPLOAD_DIR = APP_DIR / "uploads"
OUTPUT_DIR = APP_DIR / "outputs"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

app = Flask(__name__)
CORS(app)

# ---------------------------------------------------------------------------
# Model cache
# ---------------------------------------------------------------------------
MODEL_CACHE = {}

def get_model(size):
    if size not in MODEL_CACHE:
        MODEL_CACHE[size] = WhisperModel(size, device="cpu", compute_type="int8", cpu_threads=4)
    return MODEL_CACHE[size]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
LINE_RE = re.compile(r"\[(\d+(?:\.\d+)?)s -> (\d+(?:\.\d+)?)s\]\s?(.*)")

def _plain_text(raw):
    parts = []
    for line in raw.splitlines():
        m = LINE_RE.match(line.strip())
        parts.append(m.group(3) if m else line.strip())
    return " ".join(p for p in parts if p)

# ---------------------------------------------------------------------------
# Job persistence — survives server restarts
# ---------------------------------------------------------------------------
JOBS = {}

def _meta_path(job_id):
    return OUTPUT_DIR / f"{job_id}.meta.json"

def _save_job(job_id, job):
    data = {k: v for k, v in job.items() if k != "txt_path"}
    try:
        with open(_meta_path(job_id), "w", encoding="utf-8") as f:
            json.dump(data, f)
    except Exception:
        pass

def _load_job(job_id):
    path = _meta_path(job_id)
    if not path.exists():
        return None
    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        data["txt_path"] = str(OUTPUT_DIR / f"{job_id}.txt")
        if data.get("status") == "running":
            data["status"] = "error"
            data["error"]  = "Server restarted mid-job. Please upload again."
        return data
    except Exception:
        return None

# ---------------------------------------------------------------------------
# Background worker
# ---------------------------------------------------------------------------
def _worker(job_id, audio_path, model_size, lang, task, multilingual, prompt):
    job = JOBS[job_id]
    try:
        model = get_model(model_size)
        segments, info = model.transcribe(
            audio_path,
            language=lang,
            task=task,
            multilingual=multilingual,
            initial_prompt=prompt,
            vad_filter=True,
        )
        total = getattr(info, "duration", 0) or 0
        job["total"] = total
        job["lang"]  = info.language
        with open(job["txt_path"], "w", encoding="utf-8") as f:
            for seg in segments:
                f.write(f"[{seg.start:.1f}s -> {seg.end:.1f}s] {seg.text.strip()}\n")
                f.flush()
                job["done"]     = seg.end
                job["progress"] = min(seg.end / total, 0.999) if total else 0.0
        job["progress"] = 1.0
        job["status"]   = "completed"
        _save_job(job_id, job)
    except Exception as e:
        job["status"] = "error"
        job["error"]  = str(e)
        _save_job(job_id, job)
    finally:
        try:
            os.remove(audio_path)
        except Exception:
            pass

# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/", methods=["GET"])
def root():
    return render_template("index.html")

@app.route("/", methods=["HEAD"])
def root_head():
    return "", 200

@app.route("/api/health", methods=["GET", "HEAD"])
def health():
    return jsonify({"status": "ok"})


@app.route("/api/transcribe", methods=["POST"])
def start_transcribe():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file         = request.files["file"]
    model_size   = request.form.get("model", "small")
    language     = request.form.get("language", "auto-detect")
    multilingual = request.form.get("multilingual", "false").lower() == "true"
    task         = request.form.get("task", "transcribe")
    key_terms    = request.form.get("key_terms", "")

    ext        = Path(file.filename or "audio.mp3").suffix or ".audio"
    job_id     = str(uuid.uuid4())
    audio_path = str(UPLOAD_DIR / f"{job_id}{ext}")

    file.save(audio_path)

    txt_path = str(OUTPUT_DIR / f"{job_id}.txt")
    lang     = None if (language == "auto-detect" or multilingual) else language
    prompt   = key_terms.strip() or None

    JOBS[job_id] = {
        "status":   "running",
        "progress": 0.0,
        "done":     0,
        "total":    0,
        "txt_path": txt_path,
        "error":    None,
        "lang":     None,
    }

    threading.Thread(
        target=_worker,
        args=(job_id, audio_path, model_size, lang, task, multilingual, prompt),
        daemon=True,
    ).start()

    return jsonify({"job_id": job_id})


@app.route("/api/job/<job_id>", methods=["GET"])
def get_job(job_id):
    job = JOBS.get(job_id) or _load_job(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404

    transcript = ""
    if os.path.exists(job["txt_path"]):
        try:
            with open(job["txt_path"], encoding="utf-8") as f:
                transcript = f.read()
        except Exception:
            pass

    return jsonify({
        "job_id":     job_id,
        "status":     job["status"],
        "progress":   job["progress"],
        "done":       job["done"],
        "total":      job["total"],
        "lang":       job["lang"],
        "transcript": transcript,
        "error":      job.get("error"),
    })


@app.route("/api/ai/chat", methods=["POST"])
def ai_chat():
    if not GROQ_AVAILABLE:
        return jsonify({"error": "Groq not installed"}), 500
    if not GROQ_API_KEY:
        return jsonify({"error": "GROQ_API_KEY not set"}), 500

    data       = request.get_json(force=True)
    transcript = data.get("transcript", "")
    history    = data.get("history", [])
    message    = data.get("message", "")

    if not transcript.strip():
        return jsonify({"error": "No transcript provided"}), 400

    plain  = _plain_text(transcript)
    system = (
        "You are a professional document assistant. The user has transcribed an audio recording "
        "and wants help processing it into a useful document.\n\n"
        f"FULL TRANSCRIPT:\n---\n{plain[:12000]}\n---\n\n"
        "Be clear, professional, and well-structured. Use markdown: ## for headings, **bold** for key terms, - for bullets."
    )

    messages = [{"role": "system", "content": system}]
    for msg in history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": message})

    try:
        client = Groq(api_key=GROQ_API_KEY)
        resp   = client.chat.completions.create(
            model=GROQ_MODEL, messages=messages, max_tokens=4096, temperature=0.3
        )
        return jsonify({"reply": resp.choices[0].message.content})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/docx", methods=["POST"])
def generate_docx():
    if not DOCX_AVAILABLE:
        return jsonify({"error": "python-docx not installed"}), 500

    data    = request.get_json(force=True)
    content = data.get("content", "")
    title   = data.get("title", "Document")

    doc = DocxDoc()
    doc.add_heading(title, 0)

    for line in content.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
        else:
            p     = doc.add_paragraph()
            parts = re.split(r"\*\*(.*?)\*\*", line)
            for i, part in enumerate(parts):
                run      = p.add_run(part)
                run.bold = (i % 2 == 1)

    ts       = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    out_path = str(OUTPUT_DIR / f"document_{ts}.docx")
    doc.save(out_path)

    return send_file(
        out_path,
        as_attachment=True,
        download_name=f"document_{ts}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 8000)), debug=False)
